# modulo_2st.py
# Versión optimizada para despliegue en Streamlit Cloud y con mejoras de modularidad.

import pandas as pd
import numpy as np
import nltk
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.tokenize import word_tokenize
import spacy
from textblob import TextBlob
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import NMF, LatentDirichletAllocation
from docx import Document
from docx.shared import Inches
import io
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
from datetime import datetime
import re
import os
import json # Necesario para manejar la respuesta JSON del API
from concurrent.futures import ThreadPoolExecutor

# --- Configuración Inicial y Descarga de Recursos (Adaptado para Streamlit) ---

# Configuración global de gráficos
plt.rcParams["figure.figsize"] = (10, 6)
sns.set_style("whitegrid")

# Define una ruta persistente para los datos de NLTK dentro del directorio de la aplicación
NLTK_DATA_PATH = os.path.join(os.path.dirname(__file__), "nltk_data")

@st.cache_resource
def download_nltk_resources():
    """
    Descarga y configura los recursos de NLTK.
    Usa st.cache_resource para evitar descargas repetidas.
    """
    if NLTK_DATA_PATH not in nltk.data.path:
        nltk.data.path.append(NLTK_DATA_PATH)
    os.makedirs(NLTK_DATA_PATH, exist_ok=True)

    try:
        nltk.data.find('corpora/stopwords', paths=[NLTK_DATA_PATH])
    except nltk.downloader.DownloadError:
        nltk.download('stopwords', download_dir=NLTK_DATA_PATH)
    
    try:
        nltk.data.find('tokenizers/punkt', paths=[NLTK_DATA_PATH])
    except nltk.downloader.DownloadError:
        nltk.download('punkt', download_dir=NLTK_DATA_PATH)
        
# Cargar modelo de spaCy
@st.cache_resource
def load_spacy_model():
    """
    Carga el modelo de spaCy una única vez.
    """
    try:
        nlp_es = spacy.load("es_core_news_sm")
        return nlp_es
    except OSError:
        st.error("El modelo 'es_core_news_sm' de spaCy no está instalado.")
        st.info("Intentando descargar e instalar el modelo...")
        try:
            spacy.cli.download("es_core_news_sm")
            nlp_es = spacy.load("es_core_news_sm")
            return nlp_es
        except Exception as e:
            st.error(f"Error al descargar o cargar el modelo de spaCy: {e}")
            return None

# --- Funciones Auxiliares (adaptadas para Streamlit) ---

def _log_message_streamlit(message, level="info"):
    """
    Función auxiliar para mostrar mensajes en la interfaz de Streamlit.
    """
    if level == "info":
        st.info(message)
    elif level == "warning":
        st.warning(message)
    elif level == "error":
        st.error(message)
    elif level == "success":
        st.success(message)
    else:
        st.write(message) # Default para cualquier otro nivel

def _clean_text(text):
    """
    Función de pre-procesamiento de texto.
    """
    if not isinstance(text, str):
        return ""
    text = text.lower()
    text = re.sub(r'[^a-záéíóúüñ\s]', '', text) # Eliminar caracteres especiales
    text = re.sub(r'\s+', ' ', text).strip() # Eliminar espacios extra
    return text

def _get_lemmas_with_spacy(text, nlp_es):
    """
    Genera lemas para un texto usando spaCy.
    """
    if not nlp_es:
        return ""
    doc = nlp_es(text)
    lemmas = [token.lemma_ for token in doc if not token.is_stop and token.is_alpha]
    return " ".join(lemmas)

def _generate_word_frequency_chart(df_freq, title):
    """
    Genera un gráfico de barras para la frecuencia de palabras.
    """
    fig, ax = plt.subplots(figsize=(10, 6))
    sns.barplot(x='frecuencia', y='palabra', data=df_freq, ax=ax)
    ax.set_title(title)
    ax.set_xlabel("Frecuencia")
    ax.set_ylabel("Palabra")
    return fig

# --- Lógica de Análisis ---

def perform_sentiment_analysis(df, text_columns):
    """
    Realiza un análisis de sentimiento básico en las columnas de texto.
    """
    results = {}
    for col in text_columns:
        if col in df.columns:
            df[f'{col}_sentiment'] = df[col].astype(str).apply(lambda x: TextBlob(x).sentiment.polarity)
            avg_sentiment = df[f'{col}_sentiment'].mean()
            results[col] = avg_sentiment
    return results

def perform_frequency_analysis(df, text_columns):
    """
    Realiza un análisis de frecuencia de palabras y n-gramas.
    """
    download_nltk_resources()
    stop_words = set(stopwords.words('spanish'))
    results = {}
    figures = []
    
    for col in text_columns:
        if col in df.columns:
            text_series = df[col].astype(str).dropna()
            if text_series.empty:
                continue

            vectorizer = CountVectorizer(stop_words=list(stop_words), ngram_range=(1, 2))
            try:
                X = vectorizer.fit_transform(text_series)
                sum_words = X.sum(axis=0)
                words_freq = [(word, sum_words[0, idx]) for word, idx in vectorizer.vocabulary_.items()]
                words_freq = sorted(words_freq, key=lambda x: x[1], reverse=True)
                
                df_freq = pd.DataFrame(words_freq[:20], columns=['palabra', 'frecuencia'])
                results[col] = df_freq
                
                if not df_freq.empty:
                    fig = _generate_word_frequency_chart(df_freq, f"Top 20 Frecuencia de Palabras y N-gramas - Columna: {col}")
                    figures.append((col, fig))
            except ValueError:
                # Ocurre si el texto no tiene palabras después de la limpieza
                continue
                
    return results, figures

def perform_topic_modeling(df, text_columns, num_topics, topic_method, use_lemmas, nlp_es):
    """
    Realiza un modelado de temas (NMF o LDA).
    """
    results = {}
    
    for col in text_columns:
        if col in df.columns:
            text_series = df[col].astype(str).dropna()
            if text_series.empty:
                continue
                
            corpus = text_series.apply(_clean_text)
            if use_lemmas:
                corpus = corpus.apply(lambda x: _get_lemmas_with_spacy(x, nlp_es))
            
            if corpus.empty:
                continue

            vectorizer = TfidfVectorizer(max_df=0.95, min_df=2, stop_words=stopwords.words('spanish'))
            try:
                dtm = vectorizer.fit_transform(corpus)
                
                if topic_method == 'NMF':
                    model = NMF(n_components=num_topics, random_state=1)
                else:
                    model = LatentDirichletAllocation(n_components=num_topics, random_state=1)
                    
                model.fit(dtm)
                
                feature_names = vectorizer.get_feature_names_out()
                topics = []
                for topic_idx, topic in enumerate(model.components_):
                    top_words_idx = topic.argsort()[:-10 - 1:-1]
                    top_words = [feature_names[i] for i in top_words_idx]
                    topics.append(f"Tema #{topic_idx+1}: " + ", ".join(top_words))
                
                results[col] = topics
            except ValueError:
                # Ocurre si el corpus está vacío
                continue
                
    return results, None # Retorna None para el gráfico por ahora

def export_text_analysis_to_word(results_sequence, summary_text=""):
    """
    Exporta los resultados del análisis de texto a un documento de Word.
    """
    document = Document()
    document.add_heading('Informe de Análisis de Texto', 0)
    document.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if summary_text:
        document.add_paragraph(summary_text)

    for result_type, title, content in results_sequence:
        if result_type == 'text':
            document.add_heading(title, level=1)
            document.add_paragraph(content)
        elif result_type == 'image_bytes':
            document.add_heading(title, level=1)
            document.add_picture(content, width=Inches(6))
        
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf

# --- Función de Interacción con IA (Nueva y Modular) ---

async def interact_with_ai(ai_model_choice, text_for_ai, user_prompt):
    """
    Simula una interacción con un modelo de IA.
    NOTA: Para una implementación real, se debe insertar aquí la lógica
    para llamar a la API del modelo de IA (por ejemplo, Gemini).
    """
    # Placeholder de respuesta para que el código sea ejecutable
    if ai_model_choice == "Gemini":
        prompt_with_context = f"Contexto de la encuesta: {text_for_ai[:1000]}...\n\nPrompt: {user_prompt}"
        
        # Aquí iría el código para llamar a la API de Gemini
        # Ejemplo (requiere una clave de API válida):
        """
        import google.generativeai as genai
        # genai.configure(api_key="TU_API_KEY_AQUÍ")
        # model = genai.GenerativeModel('gemini-pro')
        # response = model.generate_content(prompt_with_context)
        # return response.text
        """
        
        return (
            "Este es un resultado simulado de la IA. "
            "Para obtener un resultado real, es necesario "
            "habilitar y configurar la conexión a la API "
            "de un modelo como Gemini."
        )
    else:
        return "Modelo de IA no reconocido o no configurado."

# --- FUNCIÓN PRINCIPAL DE ANÁLISIS DE TEXTO ---

def perform_text_analysis(df_input, text_columns, analysis_options):
    """
    Función principal para coordinar el análisis de texto.
    """
    # Cargar los recursos de NLTK y el modelo de spaCy
    download_nltk_resources()
    nlp_es = load_spacy_model()
    
    if nlp_es is None:
        _log_message_streamlit("❌ No se pudo cargar el modelo de spaCy, el análisis de texto será limitado.", "error")
        return None, "error", "No se pudo cargar el modelo de spaCy para lematización."

    _log_message_streamlit("🔄 Iniciando análisis de texto...", "info")
    analysis_results_sequence = []
    
    # === Análisis de Sentimiento ===
    if analysis_options.get('sentiment_analysis', False):
        sentiment_results = perform_sentiment_analysis(df_input, text_columns)
        sentiment_text_content = "Análisis de Sentimiento (polaridad promedio):\n"
        for col, avg_pol in sentiment_results.items():
            sentiment_text_content += f"- Columna '{col}': {avg_pol:.2f} (Positivo > 0, Neutro = 0, Negativo < 0)\n"
        analysis_results_sequence.append(('text', "Análisis de Sentimiento", sentiment_text_content))

    # === Frecuencia de Palabras ===
    if analysis_options.get('word_frequency', False):
        freq_results, freq_figures = perform_frequency_analysis(df_input, text_columns)
        
        freq_text_content = ""
        for col_name, res in freq_results.items():
            freq_text_content += f"📊 Frecuencia para '{col_name}':\n"
            freq_text_content += res.to_string() + "\n\n"
        analysis_results_sequence.append(('text', "Análisis de Frecuencia de Palabras y N-gramas", freq_text_content))
        
        for col_name, fig in freq_figures:
            buf = io.BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight", dpi=200)
            buf.seek(0)
            analysis_results_sequence.append(('image_bytes', f"Gráfico de Frecuencia ({col_name})", buf))
            plt.close(fig)

    # === Modelado de Temas ===
    if analysis_options.get('topic_modeling', False):
        num_topics = analysis_options.get('num_topics', 5)
        topic_method = analysis_options.get('topic_method', 'NMF')
        topic_results, _ = perform_topic_modeling(df_input, text_columns, num_topics, topic_method, use_lemmas=True, nlp_es=nlp_es)
        
        topic_text_content = "Temas identificados para las columnas seleccionadas:\n\n"
        for col, topics_list in topic_results.items():
            topic_text_content += f"**Columna: {col}**\n"
            for topic_str in topics_list:
                topic_text_content += f"- {topic_str}\n"
            topic_text_content += "\n"
        analysis_results_sequence.append(('text', "Modelado de Temas (NMF/LDA)", topic_text_content))

    # === Interacción con IA (si está habilitada) ===
    ai_model_choice = analysis_options.get('ai_model', None)
    ai_prompt = analysis_options.get('ai_prompt', None)

    if ai_model_choice and ai_prompt:
        # Concatenar todo el texto de las columnas seleccionadas
        all_text_for_ai = ""
        for col in text_columns:
            if col in df_input.columns:
                all_text_for_ai += " ".join(df_input[col].astype(str).dropna().tolist()) + " "

        all_text_for_ai = _clean_text(all_text_for_ai)
        
        if all_text_for_ai:
            _log_message_streamlit(f"Enviando {len(all_text_for_ai)} caracteres de texto a {ai_model_choice}...", "info")
            if len(all_text_for_ai) > 10000:
                _log_message_streamlit("Texto de entrada a la IA es muy largo, se truncará a 10,000 caracteres.", "warning")
                all_text_for_ai = all_text_for_ai[:10000]

            ai_response = interact_with_ai(ai_model_choice, all_text_for_ai, ai_prompt)
            
            analysis_results_sequence.append(('text', f"Respuesta de IA ({ai_model_choice})", f"**Prompt utilizado:**\n{ai_prompt}\n\n**Respuesta de la IA:**\n{ai_response}"))

    # === Exportar a Word ===
    word_document_bytes = export_text_analysis_to_word(analysis_results_sequence, "Resumen del análisis de texto...")
    
    _log_message_streamlit("✅ Análisis de texto completado.", "success")
    
    return word_document_bytes, "success", "Análisis de texto completado."