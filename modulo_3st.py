# modulo_3st.py
# Módulo para el análisis comparativo de datos de encuestas.

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import io
import numpy as np
import collections
from datetime import datetime
import streamlit as st # Importar Streamlit

# Configuración global de gráficos
plt.rcParams["figure.figsize"] = (10, 6) # Ajustado para Streamlit
sns.set_style("whitegrid")

# Mapeo de Unidades (UDs)
UD_MAPPING = {
    1: "HUGCDRN", 2: "CHUIMI", 3: "CHUC", 4: "HUNSC", 5: "GSSLZ",
    6: "AFyC GC", 7: "AFyC TFNorte", 8: "AFyC TFSur", 9: "AFyC FV",
    10: "AFyC LP", 11: "SM GC", 12: "SM TF", 13: "ObsGin CHUIMI",
    14: "ENF ObsGin", 15: "MPySP", 16: "SL", 17: "PED CHUIMI", 18: "MED Legal"
}

# --- Funciones Auxiliares (adaptadas para Streamlit) ---

def _log_message_streamlit(message, level="info"):
    """
    Función auxiliar para mostrar mensajes en la interfaz de Streamlit.
    """
    if level == "info":
        st.info(message)
    elif level == "warning":
        st.warning(message)
    elif level == "error":
        st.error(message)
    elif level == "success":
        st.success(message)
    else:
        st.write(message) # Default para cualquier otro nivel

def _generate_comparison_chart(contingency_table_percent, question_title, criterion_title, chart_type='bar'):
    """
    Genera un gráfico de barras comparativo (apilado) y lo retorna como un buffer de bytes.
    El gráfico muestra la distribución de las respuestas a una pregunta por un criterio.
    """
    fig, ax = plt.subplots(figsize=(10, 7))

    # seaborn.set_palette("viridis") # Opcional: usar una paleta de colores diferente
    contingency_table_percent.plot(kind='bar', stacked=True, ax=ax, width=0.8, colormap='viridis')

    ax.set_title(f"Distribución de '{question_title}' por '{criterion_title}'", fontsize=14)
    ax.set_xlabel(criterion_title, fontsize=12)
    ax.set_ylabel("Porcentaje (%)", fontsize=12)
    
    # Ajustar la leyenda
    ax.legend(title=f"Respuestas a {question_title}", bbox_to_anchor=(1.05, 1), loc='upper left')

    plt.xticks(rotation=45, ha='right', fontsize=10)
    plt.tight_layout(rect=[0, 0, 0.85, 1]) # Ajustar para que la leyenda no se corte

    # Guardar la figura en un buffer en memoria
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=200)
    buf.seek(0) # Resetear el puntero al principio
    plt.close(fig) # Cerrar la figura para liberar memoria
    return buf

def export_comparison_to_word(analysis_results_sequence, comparison_summary_details):
    """
    Exporta los resultados del análisis comparativo a un documento de Word.
    """
    document = Document()
    document.add_heading('Informe de Análisis Comparativo', 0)
    document.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Añadir un resumen de los criterios de comparación
    document.add_heading('Resumen de Criterios de Comparación', level=1)
    for criterion, details in comparison_summary_details.items():
        document.add_paragraph(f"- Criterio: {criterion}")
        document.add_paragraph(f"  Columnas a analizar: {', '.join(details['questions'])}")
    
    # Recorrer la secuencia de resultados y añadir al documento
    for result_type, title, content in analysis_results_sequence:
        if result_type == 'text':
            document.add_heading(title, level=2)
            document.add_paragraph(content)
        elif result_type == 'image_bytes':
            document.add_heading(title, level=2)
            document.add_picture(content, width=Inches(6))
        
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf

# --- FUNCIÓN PRINCIPAL DE ANÁLISIS COMPARATIVO ---

def perform_comparative_analysis(df_input, comparison_criteria, comparison_questions):
    """
    Función principal para coordinar el análisis comparativo.
    Genera tablas de contingencia y gráficos de distribución para cada par.
    """
    _log_message_streamlit("🔄 Iniciando análisis comparativo...", "info")
    analysis_results_sequence = []
    comparison_summary_details = {}

    for col_criterion in comparison_criteria:
        if col_criterion not in df_input.columns:
            _log_message_streamlit(f"❌ Criterio de comparación '{col_criterion}' no encontrado en el DataFrame.", "error")
            continue
        
        # Mapear los valores numéricos del criterio de comparación a nombres descriptivos
        df_input[col_criterion] = df_input[col_criterion].map(UD_MAPPING)
        crit_display_name = col_criterion
        comparison_summary_details[crit_display_name] = {'questions': comparison_questions}

        _log_message_streamlit(f"📊 Analizando por el criterio: {crit_display_name}", "info")
        
        # Iterar sobre las preguntas a comparar
        for col_question in comparison_questions:
            if col_question not in df_input.columns:
                _log_message_streamlit(f"❌ Pregunta '{col_question}' no encontrada en el DataFrame.", "warning")
                continue

            # Crear una tabla de contingencia de frecuencias
            try:
                contingency_table = pd.crosstab(
                    df_input[crit_display_name], 
                    df_input[col_question]
                )
            except Exception as e:
                _log_message_streamlit(f"Error al crear la tabla de contingencia para {col_question} por {crit_display_name}: {e}", "error")
                continue
            
            # Calcular porcentajes por fila para obtener la distribución relativa
            contingency_table_percent = contingency_table.apply(lambda r: r/r.sum() * 100, axis=1)

            # Generar el bloque de texto con los resultados
            text_result = f"\n📈 Distribución de '{col_question}' por '{crit_display_name}' (Frecuencias):\\n"
            text_result += contingency_table.to_string() + "\\n\\n"
            text_result += f"Distribución de '{col_question}' por '{crit_display_name}' (Porcentajes):\\n"
            text_result += contingency_table_percent.to_string(float_format="%.2f") + "%\\n"
            analysis_results_sequence.append(('text', f"Distribución por {crit_display_name} para {col_question}", text_result))
            
            # Generar el gráfico de barras apiladas o agrupadas
            if not contingency_table_percent.empty:
                img_buf = _generate_comparison_chart(contingency_table_percent, col_question, crit_display_name, chart_type='bar')
                analysis_results_sequence.append(('image_bytes', f"Comparación de {col_question} por {crit_display_name}", img_buf))
        
        analysis_results_sequence.append(('text', "", "-" * 80 + "\\n")) # Separador para el informe

    word_document_bytes = export_comparison_to_word(analysis_results_sequence, comparison_summary_details)
    
    _log_message_streamlit("✅ Análisis comparativo completado.", "success")
    
    return word_document_bytes, "success", "Análisis comparativo completado."
