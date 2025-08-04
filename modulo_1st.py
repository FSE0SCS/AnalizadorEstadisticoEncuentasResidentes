# modulo_1st.py

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import io # Para manejar archivos en memoria
import numpy as np
import streamlit as st # Importar Streamlit
from datetime import datetime

# Configuración global de gráficos
plt.rcParams["figure.figsize"] = (8, 5) # Tamaño fijo para todas las figuras (ancho, alto en pulgadas)
sns.set_style("whitegrid") # Estilo de fondo para gráficos

# --- Funciones Auxiliares (adaptadas para Streamlit) ---

def _log_message_streamlit(message, level="info"):
    """
    Función auxiliar para mostrar mensajes en la interfaz de Streamlit.
    """
    if level == "info":
        st.info(message)
    elif level == "warning":
        st.warning(message)
    elif level == "error":
        st.error(message)
    elif level == "success":
        st.success(message)
    else:
        st.write(message) # Default para cualquier otro nivel

def get_excel_column_letter(index):
    """
    Convierte un índice de columna (0-based) a su letra de Excel (A, B, C...).
    """
    if index < 0: return None
    result = ""
    while index >= 0:
        result = chr(65 + (index % 26)) + result
        index = index // 26 - 1
    return result

def generar_explicacion_analisis(col_name, serie_data, analysis_type, graph_style):
    """
    Genera una explicación de texto basada en los resultados del análisis.
    NOTA: Esta es una función de ejemplo, la implementación completa podría
    ser mucho más sofisticada.
    """
    explanation = f"El análisis de la columna '{col_name}' muestra una distribución de datos. "
    if analysis_type == 'quantitative':
        explanation += "Se calcularon estadísticas descriptivas como la media, la mediana y la desviación estándar para entender la tendencia central y la dispersión."
    elif analysis_type == 'categorical':
        explanation += "Se examinó la frecuencia de cada categoría para identificar las respuestas más comunes."
    
    explanation += f" El gráfico de tipo '{graph_style}' visualiza esta distribución."
    return explanation

def generar_grafico(col_name, serie_data, graph_style, analysis_type):
    """
    Genera un gráfico de matplotlib y lo retorna como un objeto de figura.
    """
    fig, ax = plt.subplots()
    
    if analysis_type == 'quantitative':
        if graph_style == 'histogram':
            sns.histplot(serie_data, kde=True, ax=ax)
            ax.set_title(f'Histograma de {col_name}')
        elif graph_style == 'boxplot':
            sns.boxplot(x=serie_data, ax=ax)
            ax.set_title(f'Diagrama de Caja de {col_name}')
        else:
            sns.histplot(serie_data, kde=True, ax=ax)
            ax.set_title(f'Histograma de {col_name}')
    elif analysis_type == 'categorical':
        if graph_style == 'bar_chart':
            value_counts = serie_data.value_counts()
            sns.barplot(x=value_counts.index, y=value_counts.values, ax=ax)
            ax.set_title(f'Gráfico de Barras de {col_name}')
            ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
        else:
            value_counts = serie_data.value_counts()
            ax.pie(value_counts.values, labels=value_counts.index, autopct='%1.1f%%', startangle=90)
            ax.axis('equal')  # Asegura que el pie chart sea un círculo
            ax.set_title(f'Gráfico de Pastel de {col_name}')
    
    plt.tight_layout()
    return fig

def export_quantitative_analysis_to_word(results_sequence, summary_text=""):
    """
    Exporta los resultados del análisis cuantitativo a un documento de Word.
    """
    document = Document()
    document.add_heading('Informe de Análisis Cuantitativo', 0)
    document.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if summary_text:
        document.add_paragraph(summary_text)

    for result_type, title, content in results_sequence:
        if result_type == 'text':
            document.add_heading(title, level=1)
            document.add_paragraph(content)
        elif result_type == 'image_bytes':
            document.add_heading(title, level=1)
            # El objeto `content` ya es un BytesIO, se puede pasar directamente
            document.add_picture(content, width=Inches(6))
        
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf

# --- FUNCIÓN PRINCIPAL DE ANÁLISIS CUANTITATIVO ---

def perform_quantitative_analysis(df_input, numeric_columns, analysis_options):
    """
    Realiza el análisis cuantitativo y genera los resultados.
    """
    _log_message_streamlit("🔄 Iniciando análisis cuantitativo...", "info")
    results_sequence = []
    
    # Generar un resumen inicial
    total_columns = len(numeric_columns)
    summary_text = f"Análisis descriptivo de {total_columns} columnas numéricas."
    
    for idx, col_name in enumerate(numeric_columns):
        if col_name in df_input.columns:
            _log_message_streamlit(f"Analizando columna '{col_name}' ({idx + 1}/{total_columns})...", "info")
            
            # Limpiar datos y convertir a tipo numérico
            serie_cleaned = pd.to_numeric(df_input[col_name], errors='coerce').dropna()
            
            if serie_cleaned.empty:
                _log_message_streamlit(f"Advertencia: La columna '{col_name}' no contiene datos numéricos válidos.", "warning")
                continue
            
            # Obtener el estilo de gráfico y tipo de análisis
            graph_style = analysis_options.get('numeric_chart_type', 'histogram')
            analysis_type = 'quantitative'
            
            # Generar texto de estadísticas
            stats_text = (
                f"Estadísticas descriptivas:\\n"
                f"Número de valores: {serie_cleaned.count():,}\\n"
                f"Media: {serie_cleaned.mean():.2f}, Mediana: {serie_cleaned.median():.2f}, Moda: {serie_cleaned.mode()[0] if not serie_cleaned.mode().empty else 'N/A'}\\n"
                f"Desviación estándar: {serie_cleaned.std():.2f}, Mín: {serie_cleaned.min():.2f}, Máx: {serie_cleaned.max():.2f}\\n"
            )
            analysis_explanation = generar_explicacion_analisis(col_name, serie_cleaned, analysis_type, graph_style)
            
            full_text_block = f"✨ Análisis de: {col_name}\\n" + stats_text + "\\n" + analysis_explanation + "\\n" + "-" * 60 + "\\n\\n"
            results_sequence.append(('text', col_name, full_text_block))

            # Generar gráfico y guardarlo en BytesIO
            fig = generar_grafico(col_name, serie_cleaned, graph_style, analysis_type)
            if fig:
                # Guardar la figura en un buffer en memoria
                buf = io.BytesIO()
                fig.savefig(buf, format="png", bbox_inches="tight", dpi=200)
                buf.seek(0) # Resetear el puntero al principio
                results_sequence.append(('image_bytes', col_name, buf))
            plt.close(fig) # Cerrar la figura para liberar memoria

    word_document_bytes = export_quantitative_analysis_to_word(results_sequence, summary_text)
    
    _log_message_streamlit("✅ Análisis cuantitativo completado.", "success")
    
    return word_document_bytes, "success", summary_text
