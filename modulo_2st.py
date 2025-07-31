# modulo_2st.py

import pandas as pd
import numpy as np
import nltk
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.tokenize import word_tokenize
import spacy
from textblob import TextBlob
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import NMF, LatentDirichletAllocation
from docx import Document
from docx.shared import Inches
import io
import matplotlib.pyplot as plt
import seaborn as sns
import google.generativeai as genai
import openai
import anthropic
# from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification # <--- COMENTAR ESTA LÍNEA
import streamlit as st
from datetime import datetime
import re
import os

# --- Configuración Inicial y Descarga de Recursos (Adaptado para Streamlit) ---

# Configuración global de gráficos
plt.rcParams["figure.figsize"] = (10, 6)
sns.set_style("whitegrid")

# Define una ruta persistente para los datos de NLTK dentro del directorio de la aplicación
NLTK_DATA_PATH = os.path.join(os.path.dirname(__file__), "nltk_data")

@st.cache_resource
def download_nltk_resources():
    if NLTK_DATA_PATH not in nltk.data.path:
        nltk.data.path.append(NLTK_DATA_PATH)
    os.makedirs(NLTK_DATA_PATH, exist_ok=True)

    try:
        nltk.data.find('corpora/stopwords', paths=[NLTK_DATA_PATH])
    except Exception as e:
        st.warning(f"Error al encontrar stopwords de NLTK, intentando descargar: {e}")
        nltk.download('stopwords', download_dir=NLTK_DATA_PATH)
    try:
        nltk.data.find('tokenizers/punkt', paths=[NLTK_DATA_PATH])
    except Exception as e:
        st.warning(f"Error al encontrar tokenizers/punkt de NLTK, intentando descargar: {e}")
        nltk.download('punkt', download_dir=NLTK_DATA_PATH)
    st.success("Recursos de NLTK (stopwords, punkt) cargados.")

@st.cache_resource
def load_spacy_model():
    try:
        nlp_es = spacy.load("es_core_news_sm")
        st.success("Modelo spaCy 'es_core_news_sm' cargado.")
        return nlp_es
    except Exception as e:
        st.error(f"Error al cargar el modelo spaCy 'es_core_news_sm': {e}. Asegúrate de que está correctamente instalado.")
        return None

download_nltk_resources()
nlp_es = load_spacy_model()

# Cargar el modelo de sentimiento RoBERTa para español
# @st.cache_resource # <--- COMENTAR ESTA LÍNEA
# def load_roberta_sentiment_model(): # <--- COMENTAR ESTA LÍNEA
#     try: # <--- COMENTAR ESTA LÍNEA
#         from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification # <--- DESCOMENTAR SI SE REACTIVA
#         tokenizer = AutoTokenizer.from_pretrained("dccuchile/bert-base-spanish-wwm-cased") # <--- COMENTAR ESTA LÍNEA
#         model = AutoModelForSequenceClassification.from_pretrained("dccuchile/bert-base-spanish-wwm-cased") # <--- COMENTAR ESTA LÍNEA
#         sentiment_pipeline = pipeline("sentiment-analysis", model=model, tokenizer=tokenizer) # <--- COMENTAR ESTA LÍNEA
#         st.success("Modelo RoBERTa para análisis de sentimiento cargado.") # <--- COMENTAR ESTA LÍNEA
#         return sentiment_pipeline # <--- COMENTAR ESTA LÍNEA
#     except Exception as e: # <--- COMENTAR ESTA LÍNEA
#         st.error(f"Error al cargar el modelo RoBERTa para sentimiento: {e}. El análisis de sentimiento avanzado no estará disponible.") # <--- COMENTAR ESTA LÍNEA
#         return None # <--- COMENTAR ESTA LÍNEA

sentiment_roberta_pipeline = None # <--- CAMBIAR A NONE PARA DESACTIVAR

# --- Funciones Auxiliares (adaptadas para Streamlit) ---

def _log_message_streamlit(message, level="info"):
    """
    Función auxiliar para mostrar mensajes en la interfaz de Streamlit.
    """
    if level == "info":
        st.info(message)
    elif level == "warning":
        st.warning(message)
    elif level == "error":
        st.error(message)
    elif level == "success":
        st.success(message)
    else:
        st.write(message) # Default

def _clean_text(text):
    """Limpia el texto: minúsculas, elimina números, signos de puntuación y espacios extra."""
    text = str(text).lower()
    text = re.sub(r'\d+', '', text) # Eliminar números
    text = re.sub(r'[^\w\s]', '', text) # Eliminar puntuación
    text = re.sub(r'\s+', ' ', text).strip() # Eliminar espacios extra
    return text

def _tokenize_and_stem(text, lang='spanish'):
    """Tokeniza, elimina stopwords y aplica stemming."""
    if not text or not isinstance(text, str):
        return []
    
    try:
        stop_words = set(stopwords.words(lang))
    except LookupError:
        nltk.download('stopwords', download_dir=NLTK_DATA_PATH)
        stop_words = set(stopwords.words(lang))

    stemmer = PorterStemmer()
    words = word_tokenize(text)
    filtered_words = [stemmer.stem(w) for w in words if w.isalpha() and w not in stop_words]
    return filtered_words

def _lemmatize_text(text, nlp_model):
    """Lemmatiza el texto usando spaCy."""
    if not text or not isinstance(text, str) or nlp_model is None:
        return []
    doc = nlp_model(text)
    lemmas = [token.lemma_ for token in doc if not token.is_stop and not token.is_punct and token.is_alpha]
    return lemmas

# --- Funciones de Análisis de PLN ---

def perform_frequency_analysis(df, text_columns, use_lemmas=False, nlp_model=None):
    """
    Realiza análisis de frecuencia de palabras y N-gramas.
    Retorna un diccionario de resultados y una lista de figuras de matplotlib.
    """
    _log_message_streamlit("Iniciando análisis de frecuencia de palabras y N-gramas...", "info")
    results = {}
    figures = []
    
    for col in text_columns:
        if col not in df.columns:
            _log_message_streamlit(f"La columna '{col}' no se encuentra en el DataFrame. Saltando.", "warning")
            continue

        corpus = df[col].astype(str).apply(_clean_text).tolist()
        corpus = [text for text in corpus if text.strip()]

        if not corpus:
            _log_message_streamlit(f"No hay texto válido en la columna '{col}' para analizar.", "warning")
            continue

        if use_lemmas and nlp_model:
            processed_text = [' '.join(_lemmatize_text(text, nlp_model)) for text in corpus]
        else:
            processed_text = [' '.join(_tokenize_and_stem(text)) for text in corpus]

        vectorizer = TfidfVectorizer(max_features=5000, ngram_range=(1, 3))
        X = vectorizer.fit_transform(processed_text)
        feature_names = vectorizer.get_feature_names_out()
        
        word_counts = X.sum(axis=0)
        frequencies = pd.DataFrame(word_counts, columns=feature_names).T.sort_values(by=0, ascending=False)
        frequencies.columns = ['Frecuencia']
        
        results[col] = frequencies.head(20)
        _log_message_streamlit(f"Análisis de frecuencia para '{col}' completado.", "success")

        fig, ax = plt.subplots(figsize=(10, 7))
        sns.barplot(x=frequencies.head(15).index, y=frequencies.head(15)['Frecuencia'], ax=ax, palette='viridis')
        ax.set_title(f"Top 15 Palabras/N-gramas en '{col}'")
        ax.set_xlabel("Palabra/N-grama")
        ax.set_ylabel("Frecuencia")
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        figures.append((col, fig))
        plt.close(fig)

    return results, figures

def perform_sentiment_analysis(df, text_columns):
    """
    Realiza análisis de sentimiento usando TextBlob y RoBERTa.
    Retorna un diccionario de resultados y una lista de figuras de matplotlib.
    """
    _log_message_streamlit("Iniciando análisis de sentimiento...", "info")
    results = {}
    figures = []

    for col in text_columns:
        if col not in df.columns:
            _log_message_streamlit(f"La columna '{col}' no se encuentra en el DataFrame. Saltando.", "warning")
            continue

        texts = df[col].astype(str).apply(_clean_text).tolist()
        texts = [text for text in texts if text.strip()]

        if not texts:
            _log_message_streamlit(f"No hay texto válido en la columna '{col}' para analizar sentimiento.", "warning")
            continue

        # TextBlob
        polarities_tb = [TextBlob(text).sentiment.polarity for text in texts]
        subjectivities_tb = [TextBlob(text).sentiment.subjectivity for text in texts]
        
        avg_polarity_tb = np.mean(polarities_tb) if polarities_tb else 0
        avg_subjectivity_tb = np.mean(subjectivities_tb) if subjectivities_tb else 0
        
        # RoBERTa (solo si el pipeline está cargado)
        roberta_sentiments = []
        if sentiment_roberta_pipeline: # <--- SOLO SE EJECUTA SI EL PIPELINE NO ES NONE
            _log_message_streamlit(f"Aplicando modelo RoBERTa para sentimiento en '{col}'...", "info")
            batch_size = 32
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i+batch_size]
                roberta_sentiments.extend(sentiment_roberta_pipeline(batch_texts))
            
            roberta_sentiment_scores = []
            for res in roberta_sentiments:
                label = res['label']
                if label == 'POS' or label == 'LABEL_2': roberta_sentiment_scores.append(res['score'])
                elif label == 'NEG' or label == 'LABEL_0': roberta_sentiment_scores.append(-res['score'])
                else: roberta_sentiment_scores.append(0)

        avg_roberta_sentiment = np.mean(roberta_sentiment_scores) if roberta_sentiment_scores else 0

        sentiment_summary = {
            "TextBlob_Polaridad_Promedio": avg_polarity_tb,
            "TextBlob_Subjetividad_Promedio": avg_subjectivity_tb,
            "RoBERTa_Sentimiento_Promedio": avg_roberta_sentiment if sentiment_roberta_pipeline else "N/A (Modelo no cargado)"
        }
        results[col] = sentiment_summary
        _log_message_streamlit(f"Análisis de sentimiento para '{col}' completado.", "success")

        fig_tb, ax_tb = plt.subplots(figsize=(8, 5))
        sns.histplot(polarities_tb, kde=True, ax=ax_tb, bins=20, color='skyblue')
        ax_tb.set_title(f"Distribución de Polaridad (TextBlob) en '{col}'")
        ax_tb.set_xlabel("Polaridad (-1 a 1)")
        ax_tb.set_ylabel("Frecuencia")
        figures.append((f"{col}_polaridad_textblob", fig_tb))
        plt.close(fig_tb)
        
        fig_sub_tb, ax_sub_tb = plt.subplots(figsize=(8, 5))
        sns.histplot(subjectivities_tb, kde=True, ax=ax_sub_tb, bins=20, color='lightgreen')
        ax_sub_tb.set_title(f"Distribución de Subjetividad (TextBlob) en '{col}'")
        ax_sub_tb.set_xlabel("Subjetividad (0 a 1)")
        ax_sub_tb.set_ylabel("Frecuencia")
        figures.append((f"{col}_subjetividad_textblob", fig_sub_tb))
        plt.close(fig_sub_tb)

        if sentiment_roberta_pipeline and roberta_sentiment_scores:
            fig_roberta, ax_roberta = plt.subplots(figsize=(8, 5))
            sns.histplot(roberta_sentiment_scores, kde=True, ax=ax_roberta, bins=20, color='coral')
            ax_roberta.set_title(f"Distribución de Sentimiento (RoBERTa) en '{col}'")
            ax_roberta.set_xlabel("Sentimiento (-1 a 1)")
            ax_roberta.set_ylabel("Frecuencia")
            figures.append((f"{col}_sentimiento_roberta", fig_roberta))
            plt.close(fig_roberta)

    return results, figures

def perform_topic_modeling(df, text_columns, num_topics=5, method='NMF', use_lemmas=False, nlp_model=None):
    """
    Realiza modelado de temas (NMF o LDA).
    Retorna un diccionario de temas y una lista de figuras de matplotlib.
    """
    _log_message_streamlit(f"Iniciando modelado de temas ({method})...", "info")
    results = {}
    figures = []

    for col in text_columns:
        if col not in df.columns:
            _log_message_streamlit(f"La columna '{col}' no se encuentra en el DataFrame. Saltando.", "warning")
            continue

        corpus = df[col].astype(str).apply(_clean_text).tolist()
        corpus = [text for text in corpus if text.strip()]

        if not corpus:
            _log_message_streamlit(f"No hay texto válido en la columna '{col}' para modelado de temas.", "warning")
            continue

        if use_lemmas and nlp_model:
            processed_text = [' '.join(_lemmatize_text(text, nlp_model)) for text in corpus]
        else:
            processed_text = [' '.join(_tokenize_and_stem(text)) for text in corpus]
        
        if method == 'LDA':
            vectorizer = CountVectorizer(max_df=0.95, min_df=2, max_features=1000, stop_words=stopwords.words('spanish'))
        else: # NMF
            vectorizer = TfidfVectorizer(max_df=0.95, min_df=2, max_features=1000, stop_words=stopwords.words('spanish'))
            
        dtm = vectorizer.fit_transform(processed_text)
        feature_names = vectorizer.get_feature_names_out()

        model = None
        if method == 'NMF':
            model = NMF(n_components=num_topics, random_state=1, init='nndsvda', max_iter=200)
        elif method == 'LDA':
            model = LatentDirichletAllocation(n_components=num_topics, random_state=1, max_iter=10)
        else:
            _log_message_streamlit(f"Método de modelado de temas '{method}' no soportado.", "error")
            continue

        model.fit(dtm)

        topics = []
        for topic_idx, topic in enumerate(model.components_):
            top_words_idx = topic.argsort()[:-10 - 1:-1]
            top_words = [feature_names[i] for i in top_words_idx]
            topics.append(f"Tema {topic_idx + 1}: {' '.join(top_words)}")
            
        results[col] = topics
        _log_message_streamlit(f"Modelado de temas para '{col}' completado.", "success")

    return results, figures

# --- Interacción con APIs de IA ---

# Inicialización de clientes de IA (usando st.secrets para las claves)
@st.cache_resource
def setup_gemini():
    api_key = st.secrets.get("GEMINI_API_KEY")
    if api_key:
        genai.configure(api_key=api_key)
        _log_message_streamlit("Google Gemini configurado.", "info")
        return genai.GenerativeModel('gemini-pro')
    else:
        _log_message_streamlit("Clave API de Gemini no encontrada en st.secrets.", "warning")
        return None

@st.cache_resource
def setup_openai():
    api_key = st.secrets.get("OPENAI_API_KEY")
    if api_key:
        openai.api_key = api_key
        _log_message_streamlit("OpenAI configurado.", "info")
        return openai.OpenAI(api_key=api_key)
    else:
        _log_message_streamlit("Clave API de OpenAI no encontrada en st.secrets.", "warning")
        return None

@st.cache_resource
def setup_anthropic():
    api_key = st.secrets.get("ANTHROPIC_API_KEY")
    if api_key:
        _log_message_streamlit("Anthropic Claude configurado.", "info")
        return anthropic.Anthropic(api_key=api_key)
    else:
        _log_message_streamlit("Clave API de Anthropic no encontrada en st.secrets.", "warning")
        return None

gemini_model = setup_gemini()
openai_client = setup_openai()
anthropic_client = setup_anthropic()


def interact_with_ai(model_choice, text_to_analyze, prompt, max_tokens=150):
    """
    Interactúa con la API de IA seleccionada.
    """
    _log_message_streamlit(f"Consultando IA ({model_choice})...", "info")
    full_prompt = f"Basado en el siguiente texto: \"{text_to_analyze}\"\n\n{prompt}"
    response_content = ""

    try:
        if model_choice == "Google Gemini" and gemini_model:
            response = gemini_model.generate_content(full_prompt)
            response_content = response.text
        elif model_choice == "OpenAI GPT" and openai_client:
            response = openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": full_prompt}],
                max_tokens=max_tokens
            )
            response_content = response.choices[0].message.content
        elif model_choice == "Anthropic Claude" and anthropic_client:
            response = anthropic_client.messages.create(
                model="claude-3-opus-20240229",
                max_tokens=max_tokens,
                messages=[{"role": "user", "content": full_prompt}],
            )
            response_content = response.content[0].text
        else:
            _log_message_streamlit(f"Modelo de IA '{model_choice}' no configurado o no soportado.", "warning")
            return "Error: Modelo de IA no configurado o no soportado."
        
        _log_message_streamlit(f"Consulta a {model_choice} completada.", "success")
        return response_content

    except Exception as e:
        _log_message_streamlit(f"Error al interactuar con {model_choice}: {e}", "error")
        return f"Error al interactuar con la IA: {e}"

# --- Exportación a Word ---

def export_qualitative_results_to_word(analysis_results_sequence):
    """
    Genera un documento Word con los resultados del análisis cualitativo.
    Recibe una secuencia de tuplas: ('text', title, content) o ('image_bytes', title, bytes_io_obj).
    Devuelve un objeto BytesIO que contiene el documento Word.
    """
    if not analysis_results_sequence:
        return None

    doc = Document()
    doc.add_heading("📝 Informe de Análisis Cualitativo y PLN", 0)
    doc.add_paragraph("Este informe detalla los análisis de procesamiento de lenguaje natural y las interacciones con modelos de inteligencia artificial.\n")

    for item_type, title, content in analysis_results_sequence:
        if item_type == 'text':
            if title:
                doc.add_heading(title, level=1) 
            doc.add_paragraph(content)
        elif item_type == 'image_bytes':
            try:
                content.seek(0)
                doc.add_picture(content, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"(Error al insertar imagen {title}: {e})")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Función Principal para Streamlit ---

def run_qualitative_analysis_streamlit(df_input, text_columns, analysis_options, ai_prompt=None, ai_model_choice=None):
    """
    Función principal para ejecutar el análisis cualitativo en Streamlit.
    """
    if df_input is None or df_input.empty:
        _log_message_streamlit("No se han cargado datos válidos para el análisis cualitativo.", "warning")
        return [], None, "NoData"

    if not text_columns:
        _log_message_streamlit("No se han seleccionado columnas de texto para el análisis.", "warning")
        return [], None, "NoTextColumns"

    analysis_results_sequence = []
    
    # === Análisis de Frecuencia ===
    if analysis_options.get('frequency', False):
        freq_results, freq_figures = perform_frequency_analysis(df_input, text_columns, use_lemmas=True, nlp_model=nlp_es)
        
        freq_text_content = "Resultados de frecuencia para las columnas seleccionadas:\n\n"
        for col, res in freq_results.items():
            freq_text_content += f"**Columna: {col}**\n"
            freq_text_content += res.to_string() + "\n\n"
        analysis_results_sequence.append(('text', "Análisis de Frecuencia de Palabras y N-gramas", freq_text_content))
        
        for col_name, fig in freq_figures:
            buf = io.BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight", dpi=200)
            buf.seek(0)
            analysis_results_sequence.append(('image_bytes', f"Gráfico de Frecuencia ({col_name})", buf))
            plt.close(fig)

    # === Análisis de Sentimiento ===
    if analysis_options.get('sentiment', False):
        sentiment_results, sentiment_figures = perform_sentiment_analysis(df_input, text_columns)
        
        sentiment_text_content = "Resultados del análisis de sentimiento para las columnas seleccionadas:\n\n"
        for col, res in sentiment_results.items():
            sentiment_text_content += f"**Columna: {col}**\n"
            sentiment_text_content += f"  Polaridad TextBlob (Promedio): {res.get('TextBlob_Polaridad_Promedio', 'N/A'):.2f}\n"
            sentiment_text_content += f"  Subjetividad TextBlob (Promedio): {res.get('TextBlob_Subjetividad_Promedio', 'N/A'):.2f}\n"
            # Asegurarse de que el valor sea numérico antes de formatear
            roberta_val = res.get('RoBERTa_Sentimiento_Promedio', 'N/A')
            if isinstance(roberta_val, (int, float)):
                sentiment_text_content += f"  Sentimiento RoBERTa (Promedio): {roberta_val:.2f}\n\n"
            else:
                sentiment_text_content += f"  Sentimiento RoBERTa (Promedio): {roberta_val}\n\n"

        analysis_results_sequence.append(('text', "Análisis de Sentimiento", sentiment_text_content))
        
        for col_name, fig in sentiment_figures:
            buf = io.BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight", dpi=200)
            buf.seek(0)
            analysis_results_sequence.append(('image_bytes', f"Gráfico de Sentimiento ({col_name})", buf))
            plt.close(fig)

    # === Modelado de Temas ===
    if analysis_options.get('topic_modeling', False):
        num_topics = analysis_options.get('num_topics', 5)
        topic_method = analysis_options.get('topic_method', 'NMF')
        topic_results, _ = perform_topic_modeling(df_input, text_columns, num_topics, topic_method, use_lemmas=True, nlp_es=nlp_es)
        
        topic_text_content = "Temas identificados para las columnas seleccionadas:\n\n"
        for col, topics_list in topic_results.items():
            topic_text_content += f"**Columna: {col}**\n"
            for topic_str in topics_list:
                topic_text_content += f"- {topic_str}\n"
            topic_text_content += "\n"
        analysis_results_sequence.append(('text', "Modelado de Temas", topic_text_content))

    # === Interacción con IA ===
    if ai_model_choice and ai_prompt:
        all_text_for_ai = ""
        for col in text_columns:
            if col in df_input.columns:
                all_text_for_ai += " ".join(df_input[col].astype(str).dropna().tolist()) + " "
        
        all_text_for_ai = _clean_text(all_text_for_ai)

        if all_text_for_ai:
            _log_message_streamlit(f"Enviando {len(all_text_for_ai)} caracteres de texto a {ai_model_choice}...", "info")
            if len(all_text_for_ai) > 10000:
                _log_message_streamlit("Texto de entrada a la IA es muy largo, se truncará a 10,000 caracteres.", "warning")
                all_text_for_ai = all_text_for_ai[:10000]

            ai_response = interact_with_ai(ai_model_choice, all_text_for_ai, ai_prompt)
            analysis_results_sequence.append(('text', f"Respuesta de IA ({ai_model_choice})", f"**Prompt utilizado:**\n{ai_prompt}\n\n**Respuesta de la IA:**\n{ai_response}"))
        else:
            _log_message_streamlit("No hay texto suficiente en las columnas seleccionadas para enviar a la IA.", "warning")

    word_doc_bytes = export_qualitative_results_to_word(analysis_results_sequence)
    _log_message_streamlit("✅ Análisis cualitativo completado.", "success")
    return analysis_results_sequence, word_doc_bytes, "Success"
